"""Генерация документации алгоритма в формате docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Стили ────────────────────────────────────────────────────────────────────

style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Calibri"
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    h.font.size = Pt(16 - (i - 1) * 2)

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def formula(text):
    """Блок формулы — моноширинный абзац с отступом."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ════════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНЫЙ ЛИСТ
# ════════════════════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MEXC S/R Signaler + Paper Trader")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Алгоритм работы программы")
r2.font.size = Pt(14)
r2.italic = True

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run("Версия 1.0  •  2025").font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. ОБЩАЯ АРХИТЕКТУРА
# ════════════════════════════════════════════════════════════════════════════

h1("1. Общая архитектура системы")

para(
    "Программа представляет собой торгового бота, который в автоматическом режиме "
    "анализирует крипторынок на бирже MEXC, обнаруживает значимые уровни поддержки "
    "и сопротивления, генерирует сигналы входа в позицию и симулирует бумажные сделки "
    "на виртуальном балансе $1 000 с плечом 10×. Все события и уведомления отправляются "
    "через Telegram-бота."
)

doc.add_paragraph()
para("Система состоит из четырёх модулей:", bold=True)
bullet("signaler.py — главный модуль: цикл анализа, Telegram-хендлеры, форматирование сообщений.")
bullet("paper_trader.py — движок бумажной торговли: управление балансом, ордерами, сделками.")
bullet("База данных — MongoDB Atlas (основное) или JSON-файлы (резерв).")
bullet("Telegram Bot API — интерфейс пользователя.")

doc.add_paragraph()

add_table(
    ["Компонент", "Технология", "Назначение"],
    [
        ["Биржевые данные", "ccxt / MEXC REST API", "OHLCV-свечи, тикеры"],
        ["Анализ данных", "pandas, numpy", "Вычисление уровней, EMA"],
        ["Telegram", "python-telegram-bot 20+", "Отправка уведомлений, кнопки"],
        ["Хранилище", "MongoDB Atlas / JSON", "Сделки, ордера, отчёты"],
        ["Планировщик", "JobQueue (PTB)", "Периодический запуск анализа"],
    ],
    col_widths=[4, 4.5, 8],
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. ЗАПУСК И ИНИЦИАЛИЗАЦИЯ
# ════════════════════════════════════════════════════════════════════════════

h1("2. Запуск и инициализация")

h2("2.1 Проверка конфигурации")
para(
    "При старте функция validate_config() проверяет наличие обязательных переменных окружения. "
    "Если хотя бы одна отсутствует — программа завершается с ошибкой."
)
doc.add_paragraph()
add_table(
    ["Переменная", "Обязательна", "Описание"],
    [
        ["MEXC_API_KEY", "Да", "API-ключ биржи MEXC"],
        ["MEXC_SECRET_KEY", "Да", "Секретный ключ биржи MEXC"],
        ["TELEGRAM_BOT_TOKEN", "Да", "Токен Telegram-бота от @BotFather"],
        ["TELEGRAM_CHAT_ID", "Да", "ID чата для отправки уведомлений"],
        ["MONGODB_URI", "Нет", "URI MongoDB Atlas; без него — JSON-файлы"],
        ["TRADING_PAIRS", "Нет", "Список пар; по умолчанию — 20 монет"],
        ["TIMEFRAME", "Нет", "Таймфрейм свечей (по умолч. 1h)"],
        ["RUN_INTERVAL_HOURS", "Нет", "Интервал анализа в часах (по умолч. 0.5)"],
        ["LEVERAGE", "Нет", "Плечо (по умолч. 10)"],
        ["SL_PERCENT", "Нет", "Стоп-лосс в % от цены входа (1.5%)"],
        ["TP_PERCENT", "Нет", "Тейк-профит в % от цены входа (3.0%)"],
        ["INITIAL_BALANCE", "Нет", "Начальный виртуальный баланс (1000$)"],
        ["TRADE_SIZE_PERCENT", "Нет", "Размер сделки, % от баланса (2%)"],
        ["MAX_OPEN_TRADES", "Нет", "Макс. одновременных позиций (5)"],
    ],
    col_widths=[5, 3, 9],
)

doc.add_paragraph()

h2("2.2 Подключение к MEXC")
para(
    "Создаётся экземпляр ccxt.mexc с ключами API. Вызов client.load_markets() загружает "
    "все доступные торговые пары биржи (~2000 инструментов). "
    "При ошибке авторизации программа завершается."
)

h2("2.3 Инициализация портфеля")
para(
    "Создаётся объект PaperPortfolio. В конструкторе:"
)
numbered("Выполняется попытка подключения к MongoDB Atlas (пинг к admin.command).")
numbered("При успехе — данные читаются из коллекций portfolio и trades.")
numbered("При неудаче или отсутствии MONGODB_URI — данные читаются из portfolio.json и trades.json.")
numbered("Восстанавливается баланс, список открытых позиций и ожидающих ордеров.")

h2("2.4 Запуск планировщика")
para(
    "Telegram Application запускается с JobQueue. Первый запуск анализа происходит через 15 секунд "
    "после старта, затем повторяется каждые RUN_INTERVAL_HOURS × 3600 секунд (по умолчанию каждые 30 минут)."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. ЦИКЛ АНАЛИЗА
# ════════════════════════════════════════════════════════════════════════════

h1("3. Цикл анализа (analysis_job)")

para(
    "Каждые 30 минут запускается функция analysis_job, которая последовательно обрабатывает "
    "все 20 монет из списка TRADING_PAIRS с паузой 2 секунды между парами (чтобы не превысить "
    "лимиты API биржи). Для каждой пары вызывается analyze_pair."
)

doc.add_paragraph()
para("Список монет по умолчанию (20 штук):", bold=True)
bullet("BTC, ETH, SOL, XRP, BNB — мажорные активы")
bullet("ADA, AVAX, DOT, ATOM, LINK — L1/DeFi")
bullet("LTC, NEAR, UNI, FIL, INJ — ликвидные альткоины")
bullet("RNDR, TON, SUI, ARB, OP — новые экосистемы")

doc.add_paragraph()
para("Структура analyze_pair (шаги выполняются по порядку):", bold=True)

add_table(
    ["Шаг", "Действие", "Результат"],
    [
        ["1", "Загрузка 250 свечей (1h) с MEXC", "DataFrame: open, high, low, close, volume"],
        ["2", "Закрытие сделок по SL/TP", "Уведомление в Telegram + обновление баланса"],
        ["3", "Проверка лимитных ордеров", "Исполнение или отмена ордеров"],
        ["4", "Поиск уровней S/R", "Список значимых уровней"],
        ["5", "Поиск сигналов входа", "Список сигналов LONG/SHORT"],
        ["6", "Сохранение отчёта в БД", "MongoDB коллекция reports / reports.json"],
        ["7", "Отправка отчёта в Telegram", "Только если есть сигналы входа"],
        ["8", "Создание лимитных ордеров", "Уведомление в Telegram о новом ордере"],
    ],
    col_widths=[1.5, 6, 9],
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. АЛГОРИТМ ПОИСКА УРОВНЕЙ
# ════════════════════════════════════════════════════════════════════════════

h1("4. Алгоритм поиска уровней поддержки и сопротивления")

h2("4.1 Нахождение локальных экстремумов")

para(
    "Алгоритм скользящего окна размером W (по умолчанию W = 8) проходит по всем свечам. "
    "Свеча с индексом i считается локальным максимумом (кандидатом в сопротивление), "
    "если её значение High строго максимально в окне [i−W, i+W]:"
)
formula("HIGH[i] = max(HIGH[i−W], ..., HIGH[i], ..., HIGH[i+W])  →  кандидат RESISTANCE")
formula("LOW[i]  = min(LOW[i−W],  ..., LOW[i],  ..., LOW[i+W])   →  кандидат SUPPORT")

para(
    "Кандидаты, у которых возраст (расстояние от конца массива свечей) меньше "
    "LEVEL_AGE_MIN_CANDLES (по умолч. 10), отфильтровываются — такие уровни считаются "
    "«незрелыми» и ненадёжными."
)

h2("4.2 Подсчёт касаний уровня")

para(
    "Для каждого кандидата подсчитывается количество свечей, которые коснулись уровня. "
    "Зона касания определяется через допуск TOLERANCE_PERCENT (по умолч. 0.8%):"
)
formula("upper = level × (1 + tolerance)")
formula("lower = level × (1 − tolerance)")
para("Свеча касается уровня, если:")
formula("(lower ≤ HIGH ≤ upper)  ИЛИ  (lower ≤ LOW ≤ upper)")

para(
    "Дополнительный фильтр объёма: если VOLUME_TOUCH_MULTIPLIER > 0, касание засчитывается "
    "только если объём свечи превышает средний объём по всей выборке, умноженный на коэффициент:"
)
formula("VOLUME[i] ≥ avg_volume × VOLUME_TOUCH_MULTIPLIER  (по умолч. 1.2)")

para(
    "Фильтр минимального расстояния между касаниями: два касания считаются разными только "
    "если между ними не менее MIN_TOUCH_SPACING свечей (по умолч. 3). "
    "Это исключает «кластерные» касания одной свечной серии."
)

para(
    "Уровень проходит фильтр, если суммарное число касаний ≥ MIN_TOUCHES (по умолч. 3)."
)

h2("4.3 Проверка ретеста (опционально)")

para(
    "Если REQUIRE_RETEST = true, уровень принимается только если после пробоя он был "
    "протестирован снова (цена вернулась к уровню). Алгоритм:"
)
numbered("Найти момент пробоя уровня (close < lower для SUPPORT или close > upper для RESISTANCE).")
numbered("После пробоя проверить, касалась ли цена зоны уровня.")
numbered("Если да — ретест подтверждён, уровень принимается.")

h2("4.4 Удаление дублей и отбор топ-N")

para(
    "Среди всех прошедших фильтр уровней удаляются дубли: два уровня одного типа "
    "(оба SUPPORT или оба RESISTANCE) считаются дублями, если их цены отличаются "
    "менее чем на TOLERANCE_PERCENT. Из дублей оставляется тот, у которого больше касаний."
)
para(
    "Оставшиеся уровни сортируются по убыванию числа касаний, и берётся "
    "топ TOP_N_LEVELS (по умолч. 5). Затем сортируются по убыванию цены для отображения."
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. АЛГОРИТМ СИГНАЛОВ ВХОДА
# ════════════════════════════════════════════════════════════════════════════

h1("5. Алгоритм генерации сигналов входа")

h2("5.1 Фильтр тренда — EMA200")

para(
    "Вычисляется экспоненциальная скользящая средняя (EMA) с периодом 200 свечей "
    "по ценам закрытия. Используется сглаживающий коэффициент α = 2 / (period + 1):"
)
formula("EMA[0] = close[0]")
formula("EMA[i] = close[i] × α  +  EMA[i−1] × (1 − α),   α = 2 / (200 + 1) ≈ 0.00995")

para("Направление тренда определяется сравнением текущей цены с EMA:")
formula("если close[-1] > EMA[-1]  →  тренд UP   (рассматриваем только LONG)")
formula("если close[-1] < EMA[-1]  →  тренд DOWN  (рассматриваем только SHORT)")

h2("5.2 Фильтр близости к уровню")

para(
    "Сигнал генерируется только если текущая цена находится вблизи уровня. "
    "Порог близости — ENTRY_PROXIMITY_PERCENT (по умолч. 0.5%):"
)
formula("distance = |current_price − level_price| / level_price × 100%")
formula("Сигнал генерируется если:  distance ≤ ENTRY_PROXIMITY_PERCENT (0.5%)")

h2("5.3 Логика направления сделки")

add_table(
    ["Тип уровня", "Тренд EMA200", "Направление сделки", "Логика"],
    [
        ["SUPPORT", "UP (цена > EMA)", "LONG 📈", "Цена отбивается от поддержки в бычьем тренде"],
        ["RESISTANCE", "DOWN (цена < EMA)", "SHORT 📉", "Цена отбивается от сопротивления в медвежьем тренде"],
        ["SUPPORT", "DOWN", "— (нет сигнала)", "Против тренда — игнорируем"],
        ["RESISTANCE", "UP", "— (нет сигнала)", "Против тренда — игнорируем"],
    ],
    col_widths=[3.5, 4, 4, 5.5],
)

doc.add_paragraph()

h2("5.4 Определение стоп-лосса и тейк-профита")

para("Точка входа — это точная цена уровня (не текущая рыночная цена). SL и TP вычисляются "
     "как процент от цены входа:")
formula("entry = level_price")
formula("")
formula("LONG:   SL = entry × (1 − SL_PERCENT / 100)    = entry × 0.985")
formula("LONG:   TP = entry × (1 + TP_PERCENT / 100)    = entry × 1.030")
formula("")
formula("SHORT:  SL = entry × (1 + SL_PERCENT / 100)    = entry × 1.015")
formula("SHORT:  TP = entry × (1 − TP_PERCENT / 100)    = entry × 0.970")
formula("")
formula("R:R = TP_PERCENT / SL_PERCENT = 3.0 / 1.5 = 1:2")

h2("5.5 Определение свечного паттерна (дополнительный фильтр)")

para(
    "Для информирования трейдера бот определяет паттерн на предпоследней завершённой свече:"
)

add_table(
    ["Паттерн", "Условие"],
    [
        ["Молот 🔨", "Тело < 30% диапазона, нижняя тень > 60%, нижняя тень > 2× верхней тени"],
        ["Падающая звезда ⭐", "Тело < 30% диапазона, верхняя тень > 60%, верхняя тень > 2× нижней"],
        ["Бычье поглощение 🕯", "Предыдущая свеча красная, текущая зелёная и полностью перекрывает её"],
        ["Медвежье поглощение 🕯", "Предыдущая свеча зелёная, текущая красная и полностью перекрывает её"],
    ],
    col_widths=[5, 12],
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. БУМАЖНАЯ ТОРГОВЛЯ
# ════════════════════════════════════════════════════════════════════════════

h1("6. Движок бумажной торговли (Paper Trader)")

h2("6.1 Жизненный цикл ордера / сделки")

para("Каждый сигнал входа проходит через четыре состояния:")
doc.add_paragraph()

add_table(
    ["Состояние", "Описание", "Переход"],
    [
        ["Сигнал", "Обнаружена точка входа у уровня", "→ Лимитный ордер"],
        ["PENDING (ожидание)", "Ордер выставлен на цену уровня, ждём касания", "→ OPEN (исполнен) или Отменён"],
        ["OPEN (открыта)", "Цена коснулась уровня, позиция открыта", "→ CLOSED (SL или TP)"],
        ["CLOSED (закрыта)", "Позиция закрыта по SL или TP, PnL зафиксирован", "— (финальное)"],
    ],
    col_widths=[4, 7, 6],
)

doc.add_paragraph()

h2("6.2 Создание лимитного ордера")

para("Проверки перед созданием ордера:")
bullet("Нет уже открытой позиции по этой паре в том же направлении.")
bullet("Нет уже ожидающего ордера по этой паре в том же направлении.")
bullet("Сумма открытых позиций и ожидающих ордеров < MAX_OPEN_TRADES (5).")

doc.add_paragraph()
para("Расчёт параметров ордера:")
formula("size_usd  = balance × TRADE_SIZE_PERCENT / 100   (маржа, $ от баланса)")
formula("notional  = size_usd × leverage                  (размер позиции в $)")
formula("risk_pct  = |entry − SL| / entry × 100           (риск в %)")
formula("reward_pct = |TP − entry| / entry × 100           (потенциал в %)")
formula("RR        = reward_pct / risk_pct                 (соотношение R:R)")

h2("6.3 Исполнение лимитного ордера")

para(
    "При каждом прогоне анализа (каждые 30 мин) для каждой пары берётся предпоследняя "
    "завершённая свеча (high и low) и проверяется, коснулась ли цена уровня:"
)
formula("LONG  исполняется если:  candle_low  ≤ entry_price")
formula("SHORT исполняется если:  candle_high ≥ entry_price")

para(
    "Каждая проверка уменьшает счётчик checks_remaining на 1. "
    "Если после PENDING_EXPIRY_CHECKS (8) проверок (≈ 4 часа) касания не произошло — "
    "ордер автоматически отменяется."
)

h2("6.4 Закрытие позиции (SL / TP)")

para("Аналогично — по данным последней свечи:")
formula("LONG  → SL срабатывает если:  candle_low  ≤ SL")
formula("LONG  → TP срабатывает если:  candle_high ≥ TP")
formula("SHORT → SL срабатывает если:  candle_high ≥ SL")
formula("SHORT → TP срабатывает если:  candle_low  ≤ TP")

h2("6.5 Расчёт PnL (прибыль/убыток)")

para("PnL рассчитывается на основе условного размера позиции (notional), а не маржи. "
     "Это отражает эффект плеча:")
formula("LONG:   PnL = (close_price − entry_price) / entry_price × notional")
formula("SHORT:  PnL = (entry_price − close_price) / entry_price × notional")
formula("")
formula("PnL_percent = PnL / size_usd × 100   (% от вложенной маржи)")
formula("")
formula("balance = balance + PnL               (обновление баланса)")

doc.add_paragraph()
para("Пример расчёта с плечом 10×:", bold=True)
add_table(
    ["Параметр", "Значение"],
    [
        ["Баланс", "$1 000"],
        ["Маржа (2%)", "$20"],
        ["Плечо", "10×"],
        ["Условный объём (notional)", "$200"],
        ["Движение цены", "+3% (TP)"],
        ["PnL", "$200 × 3% = +$6"],
        ["PnL от маржи", "+$6 / $20 × 100% = +30%"],
        ["Новый баланс", "$1 006"],
    ],
    col_widths=[7, 10],
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. ХРАНЕНИЕ ДАННЫХ
# ════════════════════════════════════════════════════════════════════════════

h1("7. Хранение данных")

h2("7.1 MongoDB Atlas (основной режим)")

para("Если задана переменная MONGODB_URI — данные хранятся в облачной базе MongoDB Atlas "
     "в БД signaler_bot:")

add_table(
    ["Коллекция", "Описание", "Ключ"],
    [
        ["portfolio", "Баланс, initial_balance, pending_orders", "_id = 'main'"],
        ["trades", "Все сделки (открытые и закрытые)", "id (UUID 8 символов)"],
        ["reports", "Последний отчёт по каждой монете", "pair (напр. 'BTC/USDT')"],
    ],
    col_widths=[4, 9, 4],
)

doc.add_paragraph()
para("Операции записи — upsert (вставка или замена). "
     "При перезапуске бот полностью восстанавливает состояние из БД.")

h2("7.2 JSON-файлы (резервный режим)")

para("Если MONGODB_URI не задан, данные сохраняются в файлы рядом со скриптом:")
add_table(
    ["Файл", "Содержимое"],
    [
        ["portfolio.json", "Баланс, initial_balance, pending_orders, updated_at"],
        ["trades.json", "Массив всех сделок"],
        ["reports.json", "Словарь pair → {text, saved_at}"],
    ],
    col_widths=[5, 12],
)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. TELEGRAM-ИНТЕРФЕЙС
# ════════════════════════════════════════════════════════════════════════════

h1("8. Telegram-интерфейс")

h2("8.1 Автоматические уведомления")

add_table(
    ["Событие", "Сообщение в чат"],
    [
        ["Найден сигнал входа", "Полный отчёт: уровни, направление, SL/TP, R:R, EMA, паттерн"],
        ["Создан лимитный ордер", "⏳ Ордер #ID: пара, направление, цена входа, SL, TP, маржа, плечо"],
        ["Ордер исполнен", "✅ Ордер исполнен: цена коснулась уровня"],
        ["Ордер отменён", "🗑 Ордер истёк: цена не дошла до уровня за 4 часа"],
        ["TP сработал", "✅ Тейк-профит: PnL в $ и %, новый баланс"],
        ["SL сработал", "❌ Стоп-лосс: PnL в $ и %, новый баланс"],
    ],
    col_widths=[5, 12],
)

doc.add_paragraph()

h2("8.2 Кнопки управления")

add_table(
    ["Кнопка", "Действие"],
    [
        ["📊 Статистика", "Баланс, изменение, winrate, кол-во сделок, лучшая/худшая сделка"],
        ["📋 Лог сделок", "Последние 10 закрытых сделок: пара, направление, PnL, причина"],
        ["📂 Позиции", "Открытые позиции с текущим нереализованным PnL + ожидающие ордера"],
        ["📈 Монеты", "Инлайн-клавиатура с 20 монетами для просмотра последнего отчёта"],
    ],
    col_widths=[4, 13],
)

doc.add_paragraph()

h2("8.3 Навигация по монетам")

para("При нажатии «📈 Монеты» бот отображает инлайн-клавиатуру с 20 кнопками (по 4 в ряд). "
     "Нажатие на монету (например, BTC) загружает из БД последний сохранённый аналитический отчёт "
     "с временной меткой обновления. Кнопка «⬅️ К монетам» возвращает к списку.")

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. БЛОК-СХЕМА ЦИКЛА
# ════════════════════════════════════════════════════════════════════════════

h1("9. Краткая блок-схема одного цикла анализа")

steps = [
    ("Старт цикла", "Каждые 30 минут, для каждой из 20 монет"),
    ("Загрузка свечей", "250 свечей 1h с MEXC через ccxt"),
    ("Проверка SL/TP", "По last candle high/low → закрытие позиций, уведомление"),
    ("Проверка ордеров", "По last candle high/low → исполнение/отмена, уведомление"),
    ("Поиск экстремумов", "Скользящее окно W=8 по high/low всей выборки"),
    ("Фильтр уровней", "Подсчёт касаний ≥ 3, объём, возраст ≥ 10 свечей, дедупликация"),
    ("EMA200", "Расчёт EMA по close; определение тренда UP/DOWN"),
    ("Фильтр близости", "Цена в радиусе 0.5% от уровня?"),
    ("Направление сделки", "SUPPORT+UP → LONG; RESISTANCE+DOWN → SHORT"),
    ("Расчёт SL/TP", "SL ±1.5%, TP ±3.0% от цены уровня"),
    ("Сохранение отчёта", "MongoDB reports / reports.json → доступно по кнопке"),
    ("Отправка сигнала", "Если есть сигналы → сообщение в Telegram"),
    ("Создание ордера", "Лимитный ордер на точной цене уровня → уведомление"),
    ("Конец пары", "Пауза 2 сек → следующая монета"),
]

table = doc.add_table(rows=len(steps), cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (step, desc) in enumerate(steps):
    cells = table.rows[i].cells
    cells[0].text = f"{i+1}. {step}"
    cells[0].paragraphs[0].runs[0].bold = True
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].text = desc
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    cells[0].width = Cm(5)
    cells[1].width = Cm(12)

    tc = cells[0]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    fill = "D6E4F0" if i % 2 == 0 else "EBF3FB"
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. ПАРАМЕТРЫ
# ════════════════════════════════════════════════════════════════════════════

h1("10. Справочник параметров")

add_table(
    ["Параметр", "По умолч.", "Описание"],
    [
        ["TIMEFRAME", "1h", "Таймфрейм свечей"],
        ["CANDLES_LIMIT", "250", "Кол-во загружаемых свечей"],
        ["EXTREMA_WINDOW", "8", "Окно поиска экстремумов (свечей)"],
        ["TOLERANCE_PERCENT", "0.8", "Допуск зоны касания уровня (%)"],
        ["MIN_TOUCHES", "3", "Мин. касаний для признания уровня"],
        ["MIN_TOUCH_SPACING", "3", "Мин. расстояние между касаниями (свечей)"],
        ["LEVEL_AGE_MIN_CANDLES", "10", "Мин. возраст уровня (свечей от конца)"],
        ["VOLUME_TOUCH_MULTIPLIER", "1.2", "Мин. объём касания (× от среднего)"],
        ["TOP_N_LEVELS", "5", "Топ уровней по числу касаний"],
        ["REQUIRE_RETEST", "false", "Требовать ретест после пробоя"],
        ["EMA_PERIOD", "200", "Период EMA для фильтра тренда"],
        ["ENTRY_PROXIMITY_PERCENT", "0.5", "Радиус близости к уровню для сигнала (%)"],
        ["SL_PERCENT", "1.5", "Стоп-лосс от цены входа (%)"],
        ["TP_PERCENT", "3.0", "Тейк-профит от цены входа (%)"],
        ["INITIAL_BALANCE", "1000", "Стартовый виртуальный баланс ($)"],
        ["TRADE_SIZE_PERCENT", "2", "Размер маржи на сделку (% от баланса)"],
        ["MAX_OPEN_TRADES", "5", "Макс. одновременных позиций"],
        ["LEVERAGE", "10", "Торговое плечо"],
        ["PENDING_EXPIRY_CHECKS", "8", "Число проверок до отмены ордера (≈4ч)"],
        ["RUN_INTERVAL_HOURS", "0.5", "Интервал между циклами анализа (часов)"],
        ["DELAY_BETWEEN_PAIRS", "2", "Пауза между парами (секунд)"],
        ["AUTO_TOP_PAIRS", "0", "0 = фиксированный список, N > 0 = топ-N по объёму"],
    ],
    col_widths=[5.5, 3, 8.5],
)

# ── ФИНАЛ ────────────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— конец документа —")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out = "/home/user/signaler_bot/algorithm.docx"
doc.save(out)
print(f"Saved: {out}")
